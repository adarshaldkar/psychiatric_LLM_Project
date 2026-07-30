"""DOCX parser — uses python-docx to extract text preserving heading structure."""
import logging
from typing import List

from app.rag.parsers.base_parser import BaseParser, ParsedDocument, ParsedPage

logger = logging.getLogger(__name__)


class DOCXParser(BaseParser):
    def can_parse(self, file_type: str) -> bool:
        return file_type.lower() == 'docx'

    def parse(self, file_path: str) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise RuntimeError("python-docx not installed")

        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            raise ValueError(f"Cannot open DOCX: {e}")

        # Extract core properties (author, title)
        props = doc.core_properties
        author = props.author or None
        title = props.title or None

        # Group paragraphs into virtual "pages" of ~50 paragraphs each
        # (DOCX has no real page concept — we approximate)
        all_paragraphs: List[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve heading markers for chunker to detect structure
            style_name = para.style.name if para.style else ''
            if 'Heading 1' in style_name:
                all_paragraphs.append(f'\n\n# {text}\n')
            elif 'Heading 2' in style_name:
                all_paragraphs.append(f'\n\n## {text}\n')
            elif 'Heading 3' in style_name:
                all_paragraphs.append(f'\n\n### {text}\n')
            else:
                all_paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    all_paragraphs.append(row_text)

        # Group into chunks of ~50 paragraphs per virtual page
        chunk_size = 50
        pages: List[ParsedPage] = []
        for i in range(0, max(1, len(all_paragraphs)), chunk_size):
            chunk = all_paragraphs[i:i + chunk_size]
            page_text = '\n'.join(chunk).strip()
            if page_text:
                pages.append(ParsedPage(
                    page_number=i // chunk_size + 1,
                    text=page_text,
                    used_ocr=False
                ))

        logger.info(f"DOCX parsed: {len(doc.paragraphs)} paragraphs → {len(pages)} virtual pages — {file_path}")

        return ParsedDocument(pages=pages, title=title, author=author, file_type='docx')
